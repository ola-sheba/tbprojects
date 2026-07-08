import io
import importlib
import json
import os
import pathlib
import re
import subprocess
import sys
from typing import Any

import pytest
from docx import Document
from fastapi.testclient import TestClient
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

APP_DIR = pathlib.Path("/app")


def _import_app():
    sys.path.insert(0, str(APP_DIR))
    module = importlib.import_module("paper_api.main")
    assert hasattr(module, "app"), "paper_api.main must expose a FastAPI object named app"
    return module.app


def _post_file(client: TestClient, name: str, data: bytes, content_type: str):
    return client.post(
        "/analyze",
        files={"file": (name, data, content_type)},
    )


def _make_large_docx() -> bytes:
    doc = Document()
    doc.add_heading("Retrieval Augmented Reasoning for Scientific Discovery", level=1)
    doc.add_paragraph(
        "Abstract. This paper studies retrieval augmented generation, scientific discovery, "
        "long context analysis, reproducibility, and academic evaluation. "
        "The core contribution is a deterministic pipeline for summarizing long papers."
    )
    base = (
        "retrieval augmented generation scientific discovery long context summarization "
        "methodology benchmark reproducibility citation evidence experiment dataset limitation "
        "evaluation neural ranking transformer abstract conclusion contribution "
    )
    # About 60,000 words, comfortably above the 50k requirement while still fast.
    paragraph = (base * 30).strip()
    for idx in range(100):
        if idx in {0, 30, 60, 90}:
            doc.add_heading(f"Section {idx}: Methods Findings Limitations", level=2)
        doc.add_paragraph(
            f"Paragraph {idx}. {paragraph} "
            "The study reports improved evidence tracing and lower hallucination risk."
        )
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _make_pdf() -> bytes:
    bio = io.BytesIO()
    c = canvas.Canvas(bio, pagesize=letter)
    text = c.beginText(72, 720)
    lines = [
        "Graph neural networks for molecular property prediction",
        "Abstract: This academic paper evaluates graph models, ablation studies, datasets, and limitations.",
        "Key finding: graph message passing improves molecular prediction accuracy with reproducible experiments.",
        "Context: the method is compared against transformer baselines and classical fingerprints.",
    ]
    for _ in range(8):
        for line in lines:
            text.textLine(line)
    c.drawText(text)
    c.showPage()
    c.save()
    return bio.getvalue()


def _assert_analysis_payload(payload: dict[str, Any], min_tokens: int):
    required = {
        "filename",
        "file_type",
        "token_count",
        "chunk_count",
        "summary",
        "key_insights",
        "context",
    }
    missing = required - set(payload)
    assert not missing, f"response missing required keys: {missing}; got {payload.keys()}"
    assert isinstance(payload["token_count"], int), "token_count must be an integer"
    assert payload["token_count"] >= min_tokens
    assert isinstance(payload["chunk_count"], int) and payload["chunk_count"] >= 1
    assert isinstance(payload["summary"], str) and len(payload["summary"].split()) >= 20
    assert len(payload["summary"]) < 6000, "summary should be concise, not the full document"
    assert isinstance(payload["key_insights"], list) and len(payload["key_insights"]) >= 3
    assert all(isinstance(item, str) and item.strip() for item in payload["key_insights"][:3])
    assert isinstance(payload["context"], dict), "context must be a structured object"


def test_project_layout_and_dependency_lockfile():
    assert (APP_DIR / "paper_api" / "main.py").exists(), "expected /app/paper_api/main.py"
    assert (APP_DIR / "paper_api" / "__init__.py").exists(), "expected /app/paper_api/__init__.py"
    assert (APP_DIR / "pyproject.toml").exists(), "expected /app/pyproject.toml"

    lockfiles = ["requirements.lock", "uv.lock", "poetry.lock", "Pipfile.lock"]
    present_lockfiles = [name for name in lockfiles if (APP_DIR / name).exists()]
    assert present_lockfiles, "provide at least one dependency lockfile with pinned versions"

    pyproject = (APP_DIR / "pyproject.toml").read_text(encoding="utf-8").lower()
    assert "fastapi" in pyproject
    assert "pypdf" in pyproject or "pdf" in pyproject
    assert "python-docx" in pyproject or "docx" in pyproject

    lock_text = "\n".join((APP_DIR / name).read_text(encoding="utf-8") for name in present_lockfiles).lower()
    assert "==" in lock_text or "version" in lock_text, "lockfile should pin dependency versions"


def test_health_endpoint():
    client = TestClient(_import_app())
    response = client.get("/health")
    assert response.status_code == 200
    payload = response.json()
    assert str(payload.get("status", "")).lower() in {"ok", "healthy", "up"}


def test_docx_long_document_upload_and_analysis():
    client = TestClient(_import_app())
    response = _post_file(
        client,
        "long_retrieval_paper.docx",
        _make_large_docx(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert response.status_code == 200, response.text
    payload = response.json()
    _assert_analysis_payload(payload, min_tokens=50_000)
    assert payload["file_type"].lower() in {"docx", ".docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    assert payload["chunk_count"] >= 3, "long documents should be chunked into multiple chunks"
    searchable = json.dumps(payload).lower()
    assert "retrieval" in searchable
    assert "reproduc" in searchable
    assert "limitation" in searchable or "method" in searchable


def test_pdf_upload_and_contextual_analysis():
    client = TestClient(_import_app())
    response = _post_file(client, "gnn_paper.pdf", _make_pdf(), "application/pdf")
    assert response.status_code == 200, response.text
    payload = response.json()
    _assert_analysis_payload(payload, min_tokens=40)
    assert payload["file_type"].lower() in {"pdf", ".pdf", "application/pdf"}
    searchable = json.dumps(payload).lower()
    assert "graph" in searchable
    assert "molecular" in searchable or "prediction" in searchable


def test_unsupported_file_type_is_rejected():
    client = TestClient(_import_app())
    response = _post_file(client, "notes.txt", b"plain text should not be accepted", "text/plain")
    assert response.status_code in {400, 415, 422}
    assert "unsupported" in response.text.lower() or "pdf" in response.text.lower() or "docx" in response.text.lower()


def test_implementation_is_local_and_not_external_llm_wrapper():
    source_files = list((APP_DIR / "paper_api").glob("**/*.py"))
    assert source_files, "no Python source files found in paper_api"
    combined = "\n".join(path.read_text(encoding="utf-8", errors="ignore").lower() for path in source_files)
    forbidden_markers = [
        "openai",
        "anthropic",
        "cohere",
        "gemini",
        "api_key",
        "authorization: bearer",
    ]
    found = [marker for marker in forbidden_markers if marker in combined]
    assert not found, f"implementation should use local parsing/NLP, found external-service markers: {found}"
